import uuid
import os
from datetime import datetime, timezone
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from app.models.document import Document, DocumentChunk, DocumentType, ProcessingStatus
from app.core.exceptions import NotFoundException, BadRequestException
from app.services.embedding_service import EmbeddingService
from loguru import logger
from typing import BinaryIO


class DocumentService:
    def __init__(self, db: AsyncSession):
        self.db = db
        self.embedding_service = EmbeddingService()

    async def upload_document(self, user_id: uuid.UUID, file: BinaryIO, filename: str) -> Document:
        ext = os.path.splitext(filename)[1].lower()
        type_map = {
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".txt": DocumentType.TXT,
        }

        file_type = type_map.get(ext)
        if not file_type:
            raise BadRequestException(f"Unsupported file type: {ext}")

        content = await self._extract_text(file, file_type)

        file.seek(0, 2)
        file_size = file.tell()
        file.seek(0)

        document = Document(
            id=uuid.uuid4(),
            user_id=user_id,
            filename=filename,
            original_filename=filename,
            file_size=file_size,
            file_type=file_type,
            content=content,
            processing_status=ProcessingStatus.PENDING,
            created_at=datetime.now(timezone.utc),
            updated_at=datetime.now(timezone.utc),
        )
        self.db.add(document)
        await self.db.flush()

        await self._process_document(document)
        return document

    async def _extract_text(self, file: BinaryIO, file_type: DocumentType) -> str:
        if file_type == DocumentType.PDF:
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file)
                return "\n".join([page.extract_text() for page in reader.pages])
            except Exception as e:
                logger.error(f"PDF extraction failed: {e}")
                return ""

        elif file_type == DocumentType.DOCX:
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file)
                return "\n".join([p.text for p in doc.paragraphs])
            except Exception as e:
                logger.error(f"DOCX extraction failed: {e}")
                return ""

        elif file_type == DocumentType.TXT:
            try:
                return file.read().decode("utf-8")
            except Exception as e:
                logger.error(f"TXT extraction failed: {e}")
                return ""

        return ""

    async def _process_document(self, document: Document):
        try:
            document.processing_status = ProcessingStatus.PROCESSING
            await self.db.flush()

            chunks = self._chunk_text(document.content)
            for i, chunk_text in enumerate(chunks):
                chunk = DocumentChunk(
                    id=uuid.uuid4(),
                    document_id=document.id,
                    chunk_index=i,
                    content=chunk_text,
                    page_number=None,
                    created_at=datetime.now(timezone.utc),
                )
                self.db.add(chunk)

                embedding_id = await self.embedding_service.add_document(
                    f"doc_{document.id}_chunk_{i}",
                    chunk_text,
                    {"document_id": str(document.id), "chunk_index": i},
                )
                chunk.embedding_id = embedding_id

            document.processing_status = ProcessingStatus.COMPLETED
            await self.db.flush()
            logger.info(f"Document processed: {document.original_filename}")

        except Exception as e:
            document.processing_status = ProcessingStatus.FAILED
            document.error_message = str(e)
            await self.db.flush()
            logger.error(f"Document processing failed: {e}")

    def _chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
        if not text:
            return []
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - overlap
        return chunks

    async def get_document(self, document_id: str, user_id: uuid.UUID) -> Document:
        result = await self.db.execute(
            select(Document).where(Document.id == document_id, Document.user_id == user_id)
        )
        doc = result.scalar_one_or_none()
        if not doc:
            raise NotFoundException("Document not found")
        return doc

    async def get_user_documents(self, user_id: uuid.UUID) -> list[Document]:
        result = await self.db.execute(
            select(Document).where(Document.user_id == user_id).order_by(Document.created_at.desc())
        )
        return result.scalars().all()

    async def delete_document(self, document_id: str, user_id: uuid.UUID):
        doc = await self.get_document(document_id, user_id)
        await self.embedding_service.delete_document(f"doc_{document_id}")
        await self.db.delete(doc)
        await self.db.flush()

    async def ask_question(self, document_id: str, user_id: uuid.UUID, question: str) -> str:
        doc = await self.get_document(document_id, user_id)

        results = await self.embedding_service.search(question, n_results=5)
        context_chunks = [r["document"] for r in results if r["metadata"].get("document_id") == document_id]

        if not context_chunks and doc.content:
            context_chunks = doc.content[:2000]

        context = "\n".join(context_chunks) if isinstance(context_chunks, list) else str(context_chunks)

        from app.services.llm_service import LLMService
        llm = LLMService()
        prompt = f"Based on the following document content, answer the question:\n\nDocument content:\n{context}\n\nQuestion: {question}"
        return await llm.generate(prompt, system_prompt="You are a document Q&A assistant. Answer based only on the provided context.")
